"""
音频转文字路由
使用 faster-whisper (tiny模型) 进行语音识别
轻量方案：模型约39MB，支持多语言
"""
from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
import os
import uuid
import json
from .utils import prepare_output_directory

router = APIRouter()

from .config import UPLOAD_DIR, OUTPUT_DIR

# faster-whisper 模型实例（延迟加载）
_whisper_model = None

# 设置 Hugging Face 镜像（解决国内网络问题）
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

def get_whisper_model():
    """获取或加载 faster-whisper 模型"""
    global _whisper_model
    if _whisper_model is None:
        try:
            from faster_whisper import WhisperModel
            _whisper_model = WhisperModel("tiny", device="cpu", compute_type="int8")
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="faster-whisper 未安装，请运行: pip install faster-whisper"
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"模型加载失败: {str(e)}。请检查网络连接或手动下载模型。"
            )
    return _whisper_model


@router.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    language: str = Form(None),
    model_size: str = Form("tiny"),
    output_format: str = Form("txt"),
    output_name: str = Form(""),
    output_path: str = Form("")
):
    """
    将音频文件转换为文字
    """
    allowed_extensions = {'.mp3', '.wav', '.m4a', '.flac', '.ogg', '.wma', '.aac', '.webm'}
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {file_ext}")
    
    file_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, f"{file_id}{file_ext}")
    
    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)
        
        model = get_whisper_model()
        
        transcribe_options = {}
        if language and language != "auto":
            transcribe_options["language"] = language
        
        segments_generator, info = model.transcribe(input_path, **transcribe_options)
        
        segments = []
        full_text_parts = []
        
        for segment in segments_generator:
            segments.append({
                "id": segment.id,
                "start": segment.start,
                "end": segment.end,
                "text": segment.text.strip()
            })
            full_text_parts.append(segment.text.strip())
        
        full_text = " ".join(full_text_parts)
        detected_language = info.language if info.language else "unknown"
        
        if detected_language == "zh" or language == "zh":
            full_text = convert_to_simplified(full_text)
            for seg in segments:
                seg["text"] = convert_to_simplified(seg["text"])
        
        output_content = ""
        output_ext = ".txt"
        
        if output_format == "srt":
            output_content = generate_srt(segments)
            output_ext = ".srt"
        elif output_format == "vtt":
            output_content = generate_vtt(segments)
            output_ext = ".vtt"
        elif output_format == "json":
            output_content = json.dumps({
                "text": full_text,
                "language": detected_language,
                "segments": segments
            }, ensure_ascii=False, indent=2)
            output_ext = ".json"
        elif output_format == "docx":
            # 生成 Word 文档
            base_name = output_name if output_name else os.path.splitext(file.filename)[0]
            output_filename = f"{base_name}.docx"
            
            # 准备输出目录
            save_dir, is_fallback = prepare_output_directory(output_path, OUTPUT_DIR)
            final_output_path = os.path.join(save_dir, output_filename)
            
            generate_docx(full_text, final_output_path)
            
            return JSONResponse({
                "success": True,
                "text": full_text,
                "language": detected_language,
                "segments": segments,
                "output_file": final_output_path,
                "output_filename": output_filename
            })
        else:
            output_content = full_text
            output_ext = ".txt"
        
        base_name = output_name if output_name else os.path.splitext(file.filename)[0]
        output_filename = f"{base_name}{output_ext}"
        
        # 准备输出目录
        save_dir, is_fallback = prepare_output_directory(output_path, OUTPUT_DIR)
        final_output_path = os.path.join(save_dir, output_filename)
        
        with open(final_output_path, "w", encoding="utf-8") as f:
            f.write(output_content)
        
        return JSONResponse({
            "success": True,
            "text": full_text,
            "language": detected_language,
            "segments": segments,
            "output_file": final_output_path,
            "output_filename": output_filename
        })
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转录失败: {str(e)}")
    finally:
        if os.path.exists(input_path):
            try:
                os.remove(input_path)
            except:
                pass


def generate_docx(text: str, output_path: str) -> None:
    """生成 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # 创建新文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 添加标题
        title = doc.add_heading('音频转文字结果', 0)
        title.alignment = 1  # 居中对齐
        
        # 添加正文
        # 按段落分割文本（以句号、问号、感叹号为分割点）
        import re
        sentences = re.split(r'[。！？.!?]', text)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence:  # 跳过空句子
                paragraph = doc.add_paragraph(sentence + '。')
                paragraph.alignment = 0  # 左对齐
        
        # 保存文档
        doc.save(output_path)
        
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx 未安装，请运行: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"生成 Word 文档失败: {str(e)}"
        )


def convert_to_simplified(text: str) -> str:
    try:
        from opencc import OpenCC
        converter = OpenCC('t2s')
        return converter.convert(text)
    except ImportError:
        return text


def generate_srt(segments: list) -> str:
    srt_content = []
    for i, seg in enumerate(segments, 1):
        start = format_timestamp_srt(seg["start"])
        end = format_timestamp_srt(seg["end"])
        srt_content.append(f"{i}\n{start} --> {end}\n{seg['text']}\n")
    return "\n".join(srt_content)


def generate_vtt(segments: list) -> str:
    vtt_content = ["WEBVTT\n"]
    for seg in segments:
        start = format_timestamp_vtt(seg["start"])
        end = format_timestamp_vtt(seg["end"])
        vtt_content.append(f"{start} --> {end}\n{seg['text']}\n")
    return "\n".join(vtt_content)


def format_timestamp_srt(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def format_timestamp_vtt(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


@router.get("/status")
async def get_status():
    try:
        from faster_whisper import WhisperModel
        installed = True
    except ImportError:
        installed = False
    
    return {
        "engine": "faster-whisper",
        "model": "tiny",
        "installed": installed,
        "supported_languages": ["zh", "en", "ja", "ko", "fr", "de", "es", "等99种语言"],
        "message": "就绪" if installed else "需要安装 faster-whisper"
    }
